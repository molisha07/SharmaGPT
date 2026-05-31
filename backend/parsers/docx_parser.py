import docx
import os
import logging

logger = logging.getLogger(__name__)

def parse_docx(file_path: str) -> str:
    """
    Extracts text from a DOCX file using python-docx.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"DOCX file not found at: {file_path}")

    try:
        doc = docx.Document(file_path)
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Also parse tables if there are any
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_content.append(" | ".join(row_text))

        return "\n".join(text_content).strip()
    except Exception as e:
        logger.error(f"Failed to parse DOCX file: {e}")
        raise RuntimeError(f"Failed to parse Word document: {e}")
